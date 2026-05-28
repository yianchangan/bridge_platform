from __future__ import annotations

import hashlib
import os
import re
import uuid
from typing import Optional

import docx
from docx.oxml.ns import nsmap as docx_nsmap

from app.models.section import ImageInfo, SectionData, TableInfo


def compute_md5(file_path: str) -> str:
    h = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _safe_filename(text: str) -> str:
    """清洗文件名: 去除非法字符, 合并连续空白"""
    cleaned = re.sub(r'[\/:*?"<>|]', '_', text)
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def scan_styles(docx_path: str) -> dict:
    """扫描文档中所有段落使用的样式及出现次数, 供用户在上传前了解文档样式分布"""
    document = docx.Document(docx_path)
    counts: dict[str, int] = {}
    for para in document.paragraphs:
        name = para.style.name.strip()
        if name:
            counts[name] = counts.get(name, 0) + 1
    return {
        "file": os.path.basename(docx_path),
        "total_paragraphs": sum(counts.values()),
        "styles": dict(sorted(counts.items(), key=lambda x: -x[1])),
    }


class WordParser:
    """基于 python-docx 的 Word 文档结构化解析器"""

    def __init__(self, doc_path: str, doc_id: str, assets_dir: str,
                 regex_config: Optional[dict] = None, max_heading_level: int = 5,
                 body_styles: Optional[list[str]] = None,
                 image_caption_style: str = "图标题",
                 table_caption_style: str = "表标题"):
        """
        Args:
            doc_path: Word 文件路径
            doc_id: 文档 UUID
            assets_dir: 资源存储根目录
            regex_config: 可选, 自定义标题正则
            max_heading_level: 最大标题层级 1-5
            body_styles: 纳入正文的段落样式名列表, 默认 ["Normal"]
            image_caption_style: 图标题样式名
            table_caption_style: 表标题样式名
        """
        self.doc_path = doc_path
        self.doc_id = doc_id
        self.doc_prefix = os.path.splitext(os.path.basename(doc_path))[0].replace(" ", "_")
        self.regex_config = regex_config or {}
        self.max_heading_level = max_heading_level
        self.body_styles = body_styles or ["Normal"]
        self.image_caption_style = image_caption_style
        self.table_caption_style = table_caption_style

        self.doc_assets_dir = os.path.join(assets_dir, doc_id)
        self.images_dir = os.path.join(self.doc_assets_dir, "images")
        self.tables_dir = os.path.join(self.doc_assets_dir, "tables")
        os.makedirs(self.images_dir, exist_ok=True)
        os.makedirs(self.tables_dir, exist_ok=True)

        self.rel_prefix = f"/storage/doc_assets/{doc_id}"

    def parse(self) -> list[SectionData]:
        """解析 Word 文档, 返回章节列表"""
        if not os.path.exists(self.doc_path):
            raise FileNotFoundError(f"文件不存在: {self.doc_path}")

        # 清理上次解析可能残留的临时图片 (命名模式: {doc_prefix}_img*.png)
        self._cleanup_temp_images()

        document = docx.Document(self.doc_path)
        sections: list[SectionData] = []
        current: Optional[dict] = None

        # 短期记忆变量
        last_seen_image_path: Optional[str] = None
        last_seen_table_caption: Optional[str] = None

        image_counter = 1
        section_counter = 0

        for element in document.element.body:
            # --- 处理段落 ---
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                para = docx.text.paragraph.Paragraph(element, document)
                style_name = para.style.name.strip()
                para_text = para.text.strip()

                # 标题
                level = self._match_heading(style_name, para_text)
                if level > 0:
                    if current is not None:
                        sections.append(self._build_section(current, section_counter))
                        section_counter += 1

                    current = {
                        "level": level,
                        "title": para_text,
                        "text": "",
                        "images": [],
                        "tables": [],
                        "tags": [],
                    }

                # 图标题: 图片在前(已暂存), 标题在后 → 绑定
                elif style_name == self.image_caption_style:
                    if last_seen_image_path and current is not None:
                        img_filename = os.path.basename(last_seen_image_path)
                        rel_path = f"{self.rel_prefix}/images/{img_filename}"
                        current["images"].append(
                            ImageInfo(caption=para_text, local_path=rel_path, original_name=img_filename).model_dump()
                        )
                        current["text"] += f"【图: {para_text} | {rel_path}】\n"
                        last_seen_image_path = None

                # 表标题: 标题在前, 表格在后 → 暂存标题
                elif style_name == self.table_caption_style:
                    last_seen_table_caption = para_text

                # 正文段落 (可配置多种样式)
                elif style_name in self.body_styles and para_text and current is not None:
                    current["text"] += para_text + "\n"

                # 检测段落中嵌入的图片实体
                if para._element.findall('.//wp:docPr', namespaces=docx.oxml.ns.nsmap):
                    for run in para.runs:
                        blips = run._element.findall('.//a:blip', namespaces=docx.oxml.ns.nsmap)
                        if blips:
                            r_embed_id = blips[0].get(
                                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                            )
                            try:
                                image_part = document.part.related_parts[r_embed_id]
                                from PIL import Image
                                import io

                                image_stream = io.BytesIO(image_part.blob)
                                with Image.open(image_stream) as img:
                                    temp_name = f"{self.doc_prefix}_img{image_counter}.png"
                                    temp_path = os.path.join(self.images_dir, temp_name)
                                    img.save(temp_path, 'PNG')
                                    image_counter += 1
                                    last_seen_image_path = temp_path
                            except Exception as e:
                                print(f"[警告] 提取图片失败: {e}")
                            break

            # --- 处理表格实体 ---
            elif isinstance(element, docx.oxml.table.CT_Tbl):
                table = docx.table.Table(element, document)
                if last_seen_table_caption and current is not None:
                    safe_cap = _safe_filename(last_seen_table_caption)
                    table_data = self._extract_table_data(table)

                    # 保存结构化 JSON
                    json_rel_path = None
                    parse_success = True

                    try:
                        import json
                        json_name = self._unique_filename(self.tables_dir, safe_cap, "json")
                        json_full = os.path.join(self.tables_dir, json_name)
                        with open(json_full, "w", encoding="utf-8") as f:
                            json.dump(
                                {"caption": last_seen_table_caption, "data": table_data},
                                f,
                                ensure_ascii=False,
                                indent=2,
                            )
                        json_rel_path = f"{self.rel_prefix}/tables/{json_name}"
                    except Exception:
                        parse_success = False

                    current["tables"].append(
                        TableInfo(
                            caption=last_seen_table_caption,
                            page_images=[],  # 后台渲染阶段填充
                            json_path=json_rel_path,
                            data=table_data if table_data else None,
                            parse_success=parse_success,
                        ).model_dump()
                    )

                    if parse_success and table_data:
                        current["text"] += f"【表: {last_seen_table_caption} | {json_rel_path}】\n"
                    else:
                        current["text"] += f"【表(图片): {last_seen_table_caption}】\n"

                    last_seen_table_caption = None

        # 最后一个章节
        if current is not None:
            sections.append(self._build_section(current, section_counter))

        return sections

    def _match_heading(self, style_name: str, para_text: str) -> int:
        """判断段落是否为标题, 返回级别 1~N, 非标题返回 0.

        优先用 regex_config 中对应级别的正则匹配段落文本,
        无正则的级别回退到 Word 内置 Heading N 样式匹配.
        """
        for lvl in range(1, self.max_heading_level + 1):
            regex = self.regex_config.get(f"level{lvl}", "").strip()
            if regex:
                if re.search(regex, para_text):
                    return lvl
            else:
                if style_name in (f"Heading {lvl}", f"heading {lvl}",
                                  f"标题 {lvl}", f"标题{lvl}"):
                    return lvl
        return 0

    @staticmethod
    def _unique_filename(directory: str, safe_basename: str, ext: str) -> str:
        """确保文件名唯一: 若 base.ext 已存在则尝试 base_2.ext, base_3.ext ..."""
        candidate = f"{safe_basename}.{ext}"
        full = os.path.join(directory, candidate)
        if not os.path.exists(full):
            return candidate
        idx = 2
        while True:
            candidate = f"{safe_basename}_{idx}.{ext}"
            full = os.path.join(directory, candidate)
            if not os.path.exists(full):
                return candidate
            idx += 1

    def _cleanup_temp_images(self):
        """清理之前解析可能残留的临时图片文件"""
        prefix = f"{self.doc_prefix}_img"
        try:
            for fname in os.listdir(self.images_dir):
                if fname.startswith(prefix) and fname.endswith(".png"):
                    try:
                        os.remove(os.path.join(self.images_dir, fname))
                    except OSError:
                        pass
        except FileNotFoundError:
            pass

    @staticmethod
    def _extract_table_data(table) -> list[list[str]]:
        """从 Table 对象提取二维文本数组"""
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        return data

    @staticmethod
    def _build_section(raw: dict, order: int) -> SectionData:
        """将解析中间数据构造成 SectionData"""
        return SectionData(
            id=str(uuid.uuid4()),
            level=raw["level"],
            title=raw["title"],
            text=raw["text"].strip(),
            images=[ImageInfo(**img) for img in raw["images"]],
            tables=[TableInfo(**tbl) for tbl in raw["tables"]],
            tags=raw.get("tags", []),
            section_order=order,
        )
